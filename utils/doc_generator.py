from docx import Document
from docx.shared import Pt
from io import BytesIO
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
import textwrap

def make_docx(country_name: str, proc_row: dict, pfm_row: dict, qa_row: dict, bottlenecks: list, recommendations: list):
    doc = Document()
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(10)

    doc.add_heading(f"{country_name} - Procurement & PFM Profile", level=1)

    # PAGE 1 - snapshot
    doc.add_heading("1. Country Snapshot & Procurement Architecture", level=2)
    doc.add_paragraph("Procurement Architecture:")
    for k,v in proc_row.items():
        doc.add_paragraph(f"• {k}: {v}", style='List Bullet')

    doc.add_heading("Quality Assurance", level=3)
    for k,v in qa_row.items():
        doc.add_paragraph(f"• {k}: {v}", style='List Bullet')

    # PAGE 2 - Bottlenecks & Recommendations
    doc.add_page_break()
    doc.add_heading("Bottlenecks & Risks", level=2)
    for b in bottlenecks:
        doc.add_paragraph(f"• {b}", style='List Bullet')

    doc.add_heading("Recommendations & Opportunities", level=2)
    for r in recommendations:
        doc.add_paragraph(f"• {r}", style='List Bullet')

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def make_pdf(country_name: str, proc_row: dict, pfm_row: dict, qa_row: dict, bottlenecks: list, recommendations: list):
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    margin = 40
    y = height - margin

    def write_heading(s):
        nonlocal y
        c.setFont("Helvetica-Bold", 14)
        c.drawString(margin, y, s)
        y -= 18

    def write_text_block(title, items):
        nonlocal y
        c.setFont("Helvetica-Bold", 11)
        c.drawString(margin, y, title)
        y -= 14
        c.setFont("Helvetica", 10)
        for it in items:
            lines = textwrap.wrap(it, 100)
            for ln in lines:
                c.drawString(margin+8, y, f"• {ln}")
                y -= 12
            y -= 4
        y -= 8

    write_heading(f"{country_name} — Procurement & PFM Profile")
    write_text_block("Procurement Architecture", [f"{k}: {v}" for k,v in proc_row.items()])
    write_text_block("PFM Snapshot", [f"{k}: {v}" for k,v in pfm_row.items()])
    write_text_block("Quality Assurance", [f"{k}: {v}" for k,v in qa_row.items()])

    c.showPage()
    y = height - margin
    write_heading("Bottlenecks & Risks")
    write_text_block("", bottlenecks)
    write_heading("Recommendations & Opportunities")
    write_text_block("", recommendations)

    c.save()
    buffer.seek(0)
    return buffer
